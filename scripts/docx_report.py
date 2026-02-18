"""
docx_report.py — Build professional .docx grading reports using python-docx.

Produces flextable-quality tables with:
  - Styled headers (EUF blue, white text)
  - Alternating row shading
  - Booktabs-style borders (horizontal only)
  - Proper font sizing and cell padding
  - Cover letter page using the university template
"""

from pathlib import Path
from typing import List, Optional, Dict, Any

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from scripts.scoring import GradingResult, DimensionScore

# ─── Brand Colors ────────────────────────────────────────────
# Approximated from the EUF template; adjust to taste.

HEADER_BG    = "1B4F72"     # Dark blue — table header background
HEADER_TEXT  = RGBColor(0xFF, 0xFF, 0xFF)
HEADING_CLR  = RGBColor(0x1B, 0x4F, 0x72)  # Dimension headings
ALT_ROW_BG   = "EBF0F5"     # Very light blue-gray — alternating rows
BORDER_DARK  = "1B4F72"     # Top/bottom rules
BORDER_LIGHT = "C8CED4"     # Inner horizontal rules
BODY_FONT    = "Calibri"
BODY_SIZE    = Pt(10.5)
SMALL_SIZE   = Pt(9.5)
HEADING_SIZE = Pt(12)


# ─── Low-level helpers ───────────────────────────────────────

def _set_cell_shading(cell, color_hex: str):
    """Apply a solid background fill to a cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _set_table_borders(table, top_sz="8", bottom_sz="8",
                       inner_h_sz="4", inner_h_color=None):
    """Apply booktabs-style borders to an entire table (no vertical lines)."""
    ih_color = inner_h_color or BORDER_LIGHT
    borders_xml = (
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top    w:val="single" w:sz="{top_sz}"    w:color="{BORDER_DARK}" w:space="0"/>'
        f'  <w:bottom w:val="single" w:sz="{bottom_sz}" w:color="{BORDER_DARK}" w:space="0"/>'
        f'  <w:left   w:val="none"   w:sz="0" w:color="auto" w:space="0"/>'
        f'  <w:right  w:val="none"   w:sz="0" w:color="auto" w:space="0"/>'
        f'  <w:insideH w:val="single" w:sz="{inner_h_sz}" w:color="{ih_color}" w:space="0"/>'
        f'  <w:insideV w:val="none"  w:sz="0" w:color="auto" w:space="0"/>'
        f'</w:tblBorders>'
    )
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    # Remove existing borders
    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)
    tblPr.append(parse_xml(borders_xml))


def _set_cell_bottom_border(cell, sz="6", color=None):
    """Add a bottom border to a single cell (for header separation)."""
    color = color or BORDER_DARK
    borders_xml = (
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:color="{color}" w:space="0"/>'
        f'</w:tcBorders>'
    )
    cell._tc.get_or_add_tcPr().append(parse_xml(borders_xml))


def _set_cell_margins(table, top=40, bottom=40, left=80, right=80):
    """Set default cell margins (padding) for the entire table, in twips."""
    margins_xml = (
        f'<w:tblCellMar {nsdecls("w")}>'
        f'  <w:top    w:w="{top}"    w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:left   w:w="{left}"   w:type="dxa"/>'
        f'  <w:right  w:w="{right}"  w:type="dxa"/>'
        f'</w:tblCellMar>'
    )
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    for existing in tblPr.findall(qn('w:tblCellMar')):
        tblPr.remove(existing)
    tblPr.append(parse_xml(margins_xml))


def _format_cell(cell, text, bold=False, font_size=None,
                 font_color=None, alignment=None, font_name=None):
    """Write formatted text into a cell, clearing existing content."""
    cell.text = ""
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = font_size or BODY_SIZE
    run.font.name = font_name or BODY_FONT
    if font_color:
        run.font.color.rgb = font_color
    # Reduce paragraph spacing inside cells
    pPr = p._p.get_or_add_pPr()
    spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:before="0" w:after="0"/>')
    for existing in pPr.findall(qn('w:spacing')):
        pPr.remove(existing)
    pPr.append(spacing)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _add_paragraph(doc, text, bold=False, font_size=None,
                   font_color=None, space_before=0, space_after=6,
                   alignment=None):
    """Add a formatted paragraph to the document."""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.bold = bold
    run.font.size = font_size or BODY_SIZE
    run.font.name = BODY_FONT
    if font_color:
        run.font.color.rgb = font_color
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    return p


def _score_to_rating(score: float, locale: dict) -> str:
    ratings = locale.get("ratings", {})
    if score >= 2.5:
        return str(ratings.get(3, "Excellent"))
    elif score >= 1.5:
        return str(ratings.get(2, "Good"))
    elif score >= 0.5:
        return str(ratings.get(1, "Adequate"))
    else:
        return str(ratings.get(0, "Insufficient"))


# ─── Table builders ──────────────────────────────────────────

def _create_styled_table(doc, headers: list, rows: list,
                         col_widths: list = None,
                         total_row: list = None,
                         font_size=None):
    """
    Create a professionally styled table.

    Parameters
    ----------
    headers : list of str
    rows    : list of list — each inner list is one row of cell values
    col_widths : list of Inches — column widths
    total_row  : optional final row rendered in bold
    font_size  : optional override for cell font size (default: BODY_SIZE)
    """
    fs = font_size or BODY_SIZE
    n_cols = len(headers)
    n_rows = len(rows) + 1 + (1 if total_row else 0)

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # ── Apply table-level formatting ──
    _set_table_borders(table)
    _set_cell_margins(table, top=20, bottom=20, left=60, right=60)

    # ── Column widths ──
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = width

    # ── Header row ──
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_shading(cell, HEADER_BG)
        _set_cell_bottom_border(cell, sz="8", color=BORDER_DARK)
        align = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
        _format_cell(cell, header_text, bold=True, font_size=fs,
                     font_color=HEADER_TEXT, alignment=align)

    # ── Body rows ──
    for row_idx, row_data in enumerate(rows):
        is_alt = row_idx % 2 == 1
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            if is_alt:
                _set_cell_shading(cell, ALT_ROW_BG)
            align = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            _format_cell(cell, str(value), font_size=fs, alignment=align)

    # ── Total / summary row ──
    if total_row:
        last_idx = len(rows) + 1
        for col_idx, value in enumerate(total_row):
            cell = table.rows[last_idx].cells[col_idx]
            _set_cell_shading(cell, HEADER_BG)
            align = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            _format_cell(cell, str(value), bold=True, font_size=fs,
                         font_color=HEADER_TEXT, alignment=align)

    return table


# ─── Page builders ───────────────────────────────────────────

PLACEHOLDER = "<Start editing here>"


def _clear_from_placeholder(doc):
    """
    Find the placeholder paragraph and remove it plus everything after it,
    preserving the template's address block, subject line, and section properties.

    The user places a line containing PLACEHOLDER (e.g., '<Start editing here>')
    in their .docx template.  Everything before it (letterhead, address, subject)
    stays untouched; everything from the placeholder onward is removed and
    replaced by the generated content.

    If no placeholder is found, falls back to appending at the end of the
    document (no content is removed).
    """
    body = doc.element.body
    placeholder_found = False

    for child in list(body):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        # Never remove section properties
        if tag == "sectPr":
            continue

        if placeholder_found:
            body.remove(child)
            continue

        # Check paragraphs for the placeholder text
        if tag == "p":
            text = child.text  # lxml .text is just the first run
            # python-docx stores text across multiple <w:r> elements,
            # so we need to concatenate all of them
            full_text = ""
            for r in child.iter(qn("w:t")):
                if r.text:
                    full_text += r.text
            if PLACEHOLDER in full_text:
                placeholder_found = True
                body.remove(child)  # remove the placeholder line itself

    if not placeholder_found:
        import click
        click.echo(f"  ⚠ Placeholder '{PLACEHOLDER}' not found in template.")
        click.echo(f"    Content will be appended at the end of the document.")
        click.echo(f"    Tip: add a line with exactly '{PLACEHOLDER}' in your .docx template.")


def _build_cover_page(doc, result: GradingResult, locale: dict,
                      lang: str, config: dict):
    """Build page 1: formal letter with summary table."""
    l = locale

    # Date (right-aligned)
    if result.date:
        _add_paragraph(doc, result.date, font_size=BODY_SIZE,
                       alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=12)

    # Title: "Gutachten Jane Doe"
    _add_paragraph(
        doc,
        f"{l['report']['title']} {result.student_name}",
        bold=True, font_size=Pt(12), space_before=4, space_after=8,
    )

    # Salutation
    salutation = l.get("cover", {}).get("salutation", "").format(
        student=result.student_name,
        title=result.thesis_title,
    )
    if salutation:
        _add_paragraph(doc, salutation, space_after=6)
    
    # Intro paragraph
    intro = l.get("cover", {}).get("intro", "").format(
        student=result.student_name,
        title=result.thesis_title,
    )
    if intro:
        _add_paragraph(doc, intro, space_after=6)

    # "Die Arbeit wird folgendermaßen bewertet:"
    assessment_lead = l.get("cover", {}).get("assessment_lead", "")
    if assessment_lead:
        _add_paragraph(doc, assessment_lead, space_after=4)

    # ── Summary table ──
    name_key = f"name_{lang}"
    headers = [
        l["report"]["dimension_label"],
        l["report"]["weight_label"],
        l["report"]["score_label"],
    ]
    rows = []
    for d in result.dimensions:
        dim_name = _dim_display_name(d, name_key)
        rows.append([dim_name, f"{d.weight*100:.0f}%", f"{d.score:.1f} / 3.0"])

    total = [
        l["report"]["overall_score"],
        "100%",
        f"{result.overall_score:.2f} / 3.0",
    ]

    _create_styled_table(
        doc, headers, rows,
        col_widths=[Inches(2.6), Inches(0.7), Inches(0.9)],
        total_row=total,
        font_size=SMALL_SIZE,
    )

    # Grade line
    grade_text = l.get("cover", {}).get("grade_line", "").format(
        percentage=f"{result.overall_percentage:.1f}%"
    )
    if grade_text:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(12)
        pf.space_after = Pt(4)
        run = p.add_run(grade_text)
        run.font.size = BODY_SIZE
        run.font.name = BODY_FONT
        run = p.add_run(" NOTE")
        run.bold = True
        run.font.size = BODY_SIZE
        run.font.name = BODY_FONT

    # Details reference
    details_note = l.get("cover", {}).get("details_note", "")
    if details_note:
        _add_paragraph(doc, details_note, space_before=4, space_after=10)

    # Closing
    closing = l.get("cover", {}).get("closing", "")
    if closing:
        _add_paragraph(doc, closing, space_after=16)

    # Signature
    _add_paragraph(doc, config.get("author", ""), space_after=0)


def _build_detail_pages(doc, result: GradingResult, locale: dict,
                        lang: str):
    """Build pages 2+: per-dimension detailed assessment."""
    l = locale
    name_key = f"name_{lang}"

    # Section heading
    _add_paragraph(
        doc,
        l["report"]["detailed"],
        bold=True, font_size=Pt(16), font_color=HEADING_CLR,
        space_before=0, space_after=12,
    )

    for d in result.dimensions:
        dim_name = _dim_display_name(d, name_key)
        rating = _score_to_rating(d.score, locale)

        # Dimension heading
        _add_paragraph(
            doc,
            f"{d.code}: {dim_name} ({d.weight*100:.0f}%)",
            bold=True, font_size=HEADING_SIZE, font_color=HEADING_CLR,
            space_before=14, space_after=2,
        )

        # Score line
        _add_paragraph(
            doc,
            f"{l['report']['score_label']}: {d.score:.1f} / 3.0 — {rating}",
            bold=True, font_size=BODY_SIZE, space_after=6,
        )

        # Sub-criteria table
        if d.sub_criteria:
            sub_headers = [
                l["report"]["sub_criteria_heading"],
                l["report"]["score_label"],
                l["report"]["rating_label"],
            ]
            sub_rows = []
            for s in d.sub_criteria:
                s_rating = _score_to_rating(s.score, locale)
                sub_rows.append([s.name, f"{s.score:.2g}", s_rating])

            _create_styled_table(
                doc, sub_headers, sub_rows,
                col_widths=[Inches(3.4), Inches(0.8), Inches(1.2)],
            )

        # Assessment comment
        if d.comment:
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_before = Pt(8)
            pf.space_after = Pt(4)
            run = p.add_run(f"{l['report']['comment_heading']}: ")
            run.bold = True
            run.font.size = BODY_SIZE
            run.font.name = BODY_FONT
            run = p.add_run(d.comment)
            run.font.size = BODY_SIZE
            run.font.name = BODY_FONT

        # Annotations
        dim_annotations = [
            a for a in result.annotations if a.get("dimension") == d.code
        ]
        if dim_annotations:
            _add_paragraph(
                doc,
                f"{l['report']['annotations_from']} ({len(dim_annotations)}):",
                bold=True, font_size=SMALL_SIZE, space_before=8, space_after=2,
            )
            for a in dim_annotations:
                page = a.get("page", "?")
                valence = a.get("valence", "")
                text = a.get("text", a.get("raw", ""))
                valence_label = locale.get("valence", {}).get(valence, "")
                prefix = f"[{l['report']['page_abbr']} {page}]"
                if valence_label:
                    prefix += f" {valence_label}"
                _add_annotation_bullet(doc, prefix, text)
        else:
            _add_paragraph(
                doc,
                l["report"]["no_annotations"],
                font_size=SMALL_SIZE, space_before=4, space_after=4,
            )


def _build_appendix(doc, result: GradingResult, locale: dict):
    """Append untagged annotations (highlights without dimension tags)."""
    if not result.untagged_annotations:
        return

    l = locale
    doc.add_page_break()

    _add_paragraph(
        doc,
        l["report"]["appendix"],
        bold=True, font_size=Pt(14), font_color=HEADING_CLR,
        space_before=0, space_after=8,
    )

    for a in result.untagged_annotations:
        page = a.get("page", "?")
        ann_type = a.get("type", "")
        text = a.get("text", a.get("raw", ""))
        prefix = f"[{l['report']['page_abbr']} {page}] ({ann_type})"
        _add_annotation_bullet(doc, prefix, text)


def _ensure_bullet_style(doc):
    """
    Ensure a 'GradingBullet' paragraph style exists with tight spacing
    and a real Word bullet list definition. Creates it once; subsequent
    calls are a no-op.
    """
    style_name = "GradingBullet"
    styles = doc.styles
    try:
        return styles[style_name]
    except KeyError:
        pass

    from docx.enum.style import WD_STYLE_TYPE

    style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles["Normal"]
    style.font.size = SMALL_SIZE
    style.font.name = BODY_FONT

    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(1)
    pf.line_spacing = 1.1
    pf.left_indent = Cm(1.0)
    pf.first_line_indent = Cm(-0.5)

    # Attach a real Word bullet via numbering XML
    # (this makes Word treat them as a proper bulleted list)
    numbering_part = doc.part.numbering_part
    numbering_elm = numbering_part._element

    # Create an abstract numbering definition with a bullet
    import random
    abs_id = random.randint(1000, 9999)
    num_id = abs_id  # reuse for simplicity

    abs_num_xml = (
        f'<w:abstractNum {nsdecls("w")} w:abstractNumId="{abs_id}">'
        f'  <w:lvl w:ilvl="0">'
        f'    <w:start w:val="1"/>'
        f'    <w:numFmt w:val="bullet"/>'
        f'    <w:lvlText w:val="\u2022"/>'
        f'    <w:lvlJc w:val="left"/>'
        f'    <w:pPr>'
        f'      <w:ind w:left="567" w:hanging="283"/>'
        f'    </w:pPr>'
        f'    <w:rPr>'
        f'      <w:rFonts w:ascii="Symbol" w:hAnsi="Symbol" w:hint="default"/>'
        f'    </w:rPr>'
        f'  </w:lvl>'
        f'</w:abstractNum>'
    )
    numbering_elm.append(parse_xml(abs_num_xml))

    num_xml = (
        f'<w:num {nsdecls("w")} w:numId="{num_id}">'
        f'  <w:abstractNumId w:val="{abs_id}"/>'
        f'</w:num>'
    )
    numbering_elm.append(parse_xml(num_xml))

    # Link the style to this numbering definition
    style_elm = style.element
    pPr = style_elm.get_or_add_pPr()
    numPr = parse_xml(
        f'<w:numPr {nsdecls("w")}>'
        f'  <w:numId w:val="{num_id}"/>'
        f'</w:numPr>'
    )
    pPr.append(numPr)

    return style


def _add_annotation_bullet(doc, prefix: str, text: str):
    """Add a compact bullet-style annotation line as a real Word list item."""
    style = _ensure_bullet_style(doc)
    p = doc.add_paragraph(style=style)

    run = p.add_run(prefix + "  ")
    run.font.size = SMALL_SIZE
    run.font.name = BODY_FONT
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    run = p.add_run(text)
    run.font.size = SMALL_SIZE
    run.font.name = BODY_FONT


def _dim_display_name(d: DimensionScore, name_key: str) -> str:
    """Get the localized name for a dimension."""
    return d.name  # already localized during scoring


# ─── Main entry point ────────────────────────────────────────

def build_docx_report(
    result: GradingResult,
    rubric: dict,
    config: dict,
    locale: dict,
    lang: str,
    output_dir: Path,
    framework_dir: Path,
) -> Path:
    """
    Build the complete .docx grading report.

    Uses the university template if available, otherwise creates
    a clean document from scratch.
    """
    ref_doc = config.get("reference_doc", "")
    template_path = framework_dir / ref_doc if ref_doc else None

    if template_path and template_path.exists():
        doc = Document(str(template_path))
        _clear_from_placeholder(doc)
    else:
        doc = Document()
        # Set default font for new documents
        style = doc.styles["Normal"]
        style.font.name = BODY_FONT
        style.font.size = BODY_SIZE

    # Page 1: cover letter
    _build_cover_page(doc, result, locale, lang, config)

    # Page break → detailed assessment
    doc.add_page_break()
    _build_detail_pages(doc, result, locale, lang)

    # Appendix
    _build_appendix(doc, result, locale)

    # Footer note
    # _add_paragraph(
    #     doc,
    #     locale["report"]["generated_note"],
    #     font_size=Pt(8),
    #     font_color=RGBColor(0x99, 0x99, 0x99),
    #     space_before=20,
    #     alignment=WD_ALIGN_PARAGRAPH.CENTER,
    # )

    # Save
    output_path = output_dir / "report.docx"
    doc.save(str(output_path))
    return output_path
